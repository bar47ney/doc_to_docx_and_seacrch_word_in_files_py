import PySimpleGUI as sg
import os
import docx
from docx import Document
import time
from threading import Thread
import sys



def main():
    mylist = [1,2,3,4,5,6,7,8]

    progressbar = [
        [sg.ProgressBar(len(mylist), orientation='h', size=(48, 10), key='progressbar')]
    ]
    
    _STATE = [False]
    
    class Thread1(Thread):
        def run(self):        
        # def tt1():    
            progress_bar.UpdateBar(0)
            a = 1
            while a==1:
                for i,item in enumerate(mylist):
                    # print(item)
                    time.sleep(1)
                    progress_bar.UpdateBar(i + 1)
                    # print(f'-----------------------------{_STATE}')
                    if not _STATE[0]:                
                        progress_bar.UpdateBar(8)
                        break
                progress_bar.UpdateBar(0)
                if not _STATE[0]:                
                    progress_bar.UpdateBar(8)
                    print('-------------Поиск окончен-------------')
                    break
                
    class Thread2(Thread):
        def run(self):
    # def tt2():
            _STATE[0] = True            
            window['-OUTPUT-'].update('')
            text = values['-WORD-'].lower()
            print(f'Ищу - {text}')
            paths = []
            folder = os.getcwd()
            for root, dirs, files in os.walk(folder):
                for file in files:
                    if file.endswith('docx') and not file.startswith('~'):
                        paths.append(os.path.join(root, file))
            counter = 0
            # print(f'Введите слово')
            # text = input()
            mydoc = docx.Document()
            for path in paths:  
                doc = docx.Document(path)
                properties = doc.core_properties
                for paragraph in doc.paragraphs:
                    if text in paragraph.text.lower():
                        counter += 1    
                        file_name_with_extension = path.split("/")[-1]
                        file_name = file_name_with_extension.split(".")[0]
                        url = f'{file_name}_диклофенак.docx'  
                        mydoc.add_paragraph(file_name)   
                        print(f'{file_name}_диклофенак')
                        print(counter)
                        break
            mydoc.add_paragraph(f'Найдено совпадений - {counter}')       
            mydoc.save("result.docx")
            _STATE[0] = False
    
    
    layout = [
        [sg.Text('Введите слово'), sg.InputText(key='-WORD-', do_not_clear=False)],
        [sg.Frame('Progress',layout= progressbar)],
        [sg.Output(size=(88, 20), key='-OUTPUT-')],
        [sg.Submit("Поиск"), sg.Cancel("Закрыть")]
    ]

    window = sg.Window('Поиск по словам', layout)
    progress_bar = window['progressbar']
    while True:                             # The Event Loop
        event, values = window.read()
        # print(event, values) #debug
        if event in (None, 'Exit', 'Закрыть'):
            break
            sys.exit()
        if event == 'Поиск':        
            # logic(text)
            t1 = Thread1()
            t1.start()
            t2 = Thread2()
            t2.start()
            # t1 = Thread(target=tt1(), daemon=True)
            # t1.setDaemon(True)
            # t1.isDaemon()
            # t1.start()
            # t2 = Thread(target=tt2(), daemon=True)
            # t2.setDaemon(True)
            # t2.isDaemon()
            # t2.start()
    window.close()   



if __name__ == "__main__":
    main()


# find





# doc to docx

# from glob import glob
# import re
# import os
# import win32com.client as win32
# from win32com.client import constants

# # Create list of paths to .doc files
# paths = glob('C:\\Users\\Сергей\\Downloads\\pyth\\2022 — копия\\*.doc', recursive=True)

# def save_as_docx(path):
#     # Opening MS Word
#     word = win32.gencache.EnsureDispatch('Word.Application')
#     doc = word.Documents.Open(path)
#     doc.Activate ()

#     # Rename path with .docx
#     new_file_abs = os.path.abspath(path)
#     new_file_abs = re.sub(r'\.\w+$', '.docx', new_file_abs)

#     # Save and Close
#     word.ActiveDocument.SaveAs(
#         new_file_abs, FileFormat=constants.wdFormatXMLDocument
#     )
#     doc.Close(False)

# for path in paths:
#     save_as_docx(path)